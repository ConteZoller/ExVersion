import argparse
import os
import sys
import json
from fpdf import FPDF
import fitz
import markdown2
import pandas as pd
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx2pdf import convert as docx_to_pdf
from odf.opendocument import load, OpenDocumentText
from odf import text as odf_text
from bs4 import BeautifulSoup
from io import StringIO


def convert_txt_to_pdf(input_file, output_file):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    with open(input_file, 'r', encoding='utf-8') as file:
        for line in file:
            pdf.cell(200, 10, txt=line, ln=True)

    pdf.output(output_file)

def convert_txt_to_csv(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as input_file:
        lines = input_file.readlines()

    with open(output_file, 'w', encoding='utf-8', newline='') as output_file:
        for line in lines:
            output_file.write(line)

def convert_txt_to_json(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as input_file:
        lines = input_file.readlines()

    data = {"lines": lines}

    with open(output_file, 'w', encoding='utf-8') as output_file:
        json.dump(data, output_file, indent=2)

def convert_txt_to_md(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as input_file:
        text = input_file.read()

    md_content = markdown2.markdown(text)

    with open(output_file, 'w', encoding='utf-8') as output_file:
        output_file.write(md_content)

def convert_txt_to_xlsx(input_file, output_file):
    df = pd.read_csv(input_file, delimiter='\t', header=None)
    df.to_excel(output_file, index=False, header=False)

def convert_txt_to_doc(input_file, output_file):
    doc = Document()
    
    with open(input_file, 'r', encoding='utf-8') as file:
        for line in file:
            doc.add_paragraph(line.strip())

    doc.save(output_file)

def convert_txt_to_docx(input_file, output_file):
    doc = Document()

    with open(input_file, 'r', encoding='utf-8') as file:
        for line in file:
            doc.add_paragraph(line.strip())

    doc.save(output_file)

def convert_txt_to_rtf(input_file, output_file):
    doc = Document()
    
    with open(input_file, 'r', encoding='utf-8') as file:
        for line in file:
            doc.add_paragraph(line.strip())

    doc.save(output_file, format='rtf')

def convert_txt_to_odt(input_file, output_file):
    doc = OpenDocumentText()

    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()

    para = odf_text.P(text=content)
    doc.text.addElement(para)

    doc.save(output_file)

def convert_txt_to_html(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as input_file:
        text = input_file.read()

    soup = BeautifulSoup(text, 'html.parser')
    formatted_text = soup.prettify()

    with open(output_file, 'w', encoding='utf-8') as output_file:
        output_file.write(formatted_text)

##########################################################################################################

def convert_pdf_to_txt(input_file, output_file):
    doc = fitz.open(input_file)
    text = ""

    for page_number in range(doc.page_count):
        page = doc[page_number]
        text += page.get_text()

    with open(output_file, 'w', encoding='utf-8') as output_file:
        output_file.write(text)

def convert_pdf_to_csv(input_file, output_file):
    doc = fitz.open(input_file)
    text = ""

    for page_number in range(doc.page_count):
        page = doc[page_number]
        text += page.get_text()

    # Assuming that the PDF content is tabular data
    df = pd.read_csv(StringIO(text), sep="\t", header=None)
    df.to_csv(output_file, sep='\t', index=False, header=False)

def convert_pdf_to_json(input_file, output_file):
    doc = fitz.open(input_file)
    text = ""

    for page_number in range(doc.page_count):
        page = doc[page_number]
        text += page.get_text()

    data = {"lines": text.splitlines()}

    with open(output_file, 'w', encoding='utf-8') as json_file:
        json.dump(data, json_file, indent=2)

def convert_pdf_to_md(input_file, output_file):
    doc = fitz.open(input_file)
    text = ""

    for page_number in range(doc.page_count):
        page = doc[page_number]
        text += page.get_text()

    md_content = markdown2.markdown(text)

    with open(output_file, 'w', encoding='utf-8') as md_file:
        md_file.write(md_content)


def convert_pdf_to_xlsx(input_file, output_file):
    # Implementation for converting PDF to XLSX using openpyxl
    df_list = []
    doc = fitz.open(input_file)

    for page_number in range(doc.page_count):
        page = doc[page_number]
        text = page.get_text()
        df_list.append([text])

    df = pd.DataFrame(df_list, columns=['Text'])
    df.to_excel(output_file, index=False)

def convert_pdf_to_doc(input_file, output_file):
    # Implementation for converting PDF to DOC using python-docx
    doc = Document()
    
    # Iterate through each page of the PDF and add text to the document
    pdf_document = fitz.open(input_file)
    for page_number in range(pdf_document.page_count):
        page = pdf_document[page_number]
        text = page.get_text()
        
        # Add the text as a paragraph to the document
        paragraph = doc.add_paragraph(text)
        
        # Adjust font size and alignment if needed
        for run in paragraph.runs:
            run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(output_file)

def convert_pdf_to_docx(input_file, output_file):
    # Implementation for converting PDF to DOCX using python-docx and PyMuPDF
    doc = Document()

    # Iterate through each page of the PDF and add text to the document
    pdf_document = fitz.open(input_file)
    for page_number in range(pdf_document.page_count):
        page = pdf_document[page_number]
        text = page.get_text()
        
        # Add the text as a paragraph to the document
        paragraph = doc.add_paragraph(text)
        
        # Adjust font size and alignment if needed
        for run in paragraph.runs:
            run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(output_file)

def convert_pdf_to_xls(input_file, output_file):
    # Implementation for converting PDF to XLS using pandas
    df_list = []
    doc = fitz.open(input_file)

    for page_number in range(doc.page_count):
        page = doc[page_number]
        text = page.get_text()
        df_list.append([text])

    df = pd.DataFrame(df_list, columns=['Text'])
    df.to_excel(output_file, index=False)

def convert_pdf_to_rtf(input_file, output_file):
    # Implementation for converting PDF to RTF using reportlab
    doc = fitz.open(input_file)
    text = ""

    for page_number in range(doc.page_count):
        page = doc[page_number]
        text += page.get_text()

    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(text)

def convert_pdf_to_odt(input_file, output_file):
    # Implementation for converting PDF to ODT using python-docx
    doc = Document()

    # Iterate through each page of the PDF and add text to the document
    pdf_document = fitz.open(input_file)
    for page_number in range(pdf_document.page_count):
        page = pdf_document[page_number]
        text = page.get_text()
        
        # Add the text as a paragraph to the document
        paragraph = doc.add_paragraph(text)
        
        # Adjust font size and alignment if needed
        for run in paragraph.runs:
            run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(output_file)

def convert_pdf_to_html(input_file, output_file):
    # Implementation for converting PDF to HTML using BeautifulSoup
    doc = fitz.open(input_file)
    text = ""

    for page_number in range(doc.page_count):
        page = doc[page_number]
        text += page.get_text()

    soup = BeautifulSoup(text, 'html.parser')
    formatted_text = soup.prettify()

    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(formatted_text)

##########################################################################################################

def convert_csv_to_txt(input_file, output_file):
    df = pd.read_csv(input_file, header=None, delimiter='\t')
    df.to_csv(output_file, sep='\t', index=False, header=False)

##########################################################################################################

def convert_json_to_txt(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as json_file:
        data = json.load(json_file)
        lines = data.get("lines", [])
    
    with open(output_file, 'w', encoding='utf-8') as txt_file:
        for line in lines:
            txt_file.write(line)

##########################################################################################################

def convert_md_to_txt(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as md_file:
        md_content = md_file.read()
    
    with open(output_file, 'w', encoding='utf-8') as txt_file:
        txt_file.write(md_content)

##########################################################################################################

def main():
    parser = argparse.ArgumentParser(description='Converti un file in vari formati.')
    parser.add_argument('file_da_convertire', help='Il percorso del file da convertire')
    parser.add_argument('formato_di_conversione', choices=['txt', 'pdf', 'csv', 'json', 'md', 'xlsx', 'doc', 'docx', 'rtf', 'odt', 'html'], help='Il formato di conversione desiderato')
    
    args = parser.parse_args()

    if not os.path.isfile(args.file_da_convertire):
        print(f"Errore: Il file '{args.file_da_convertire}' non esiste.")
        sys.exit(1)

    base_name, extension = os.path.splitext(args.file_da_convertire)
    output_file = f"{base_name}.{args.formato_di_conversione}"

    count = 1
    while os.path.exists(output_file):
        output_file = f"{base_name}({count}).{args.formato_di_conversione}"
        count += 1

    if extension.lower() == '.txt':
        if args.formato_di_conversione == 'pdf':
            convert_txt_to_pdf(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file PDF è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'csv':
            convert_txt_to_csv(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file CSV è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'json':
            convert_txt_to_json(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file JSON è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'md':
            convert_txt_to_md(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file MD è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'xlsx':
            convert_txt_to_xlsx(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file XLSX è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'doc':
            convert_txt_to_doc(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file DOC è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'docx':
            convert_txt_to_docx(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file DOCX è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'rtf':
            convert_txt_to_rtf(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file RTF è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'odt':
            convert_txt_to_odt(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file ODT è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'html':
            convert_txt_to_html(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file HTML è stato salvato come '{output_file}'.")
        else:
            print(f"Errore: Il formato di conversione '{args.formato_di_conversione}' non è supportato per i file di testo. Scegli un formato tra 'pdf', 'csv', 'json', 'md', 'xlsx', 'doc', 'docx', 'rtf', 'odt', 'html'.")
            sys.exit(1)

    elif extension.lower() == '.pdf':
        if args.formato_di_conversione == 'txt':
            convert_pdf_to_txt(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file di testo è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'csv':
            convert_pdf_to_csv(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file CSV è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'json':
            convert_pdf_to_json(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file JSON è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'md':
            convert_pdf_to_md(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file MD è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'xlsx':
            convert_pdf_to_xlsx(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file XLSX è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'doc':
            convert_pdf_to_doc(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file DOC è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'docx':
            convert_pdf_to_docx(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file DOCX è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'rtf':
            convert_pdf_to_rtf(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file RTF è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'odt':
            convert_pdf_to_odt(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file ODT è stato salvato come '{output_file}'.")
        elif args.formato_di_conversione == 'html':
            convert_pdf_to_html(args.file_da_convertire, output_file)
            print(f"Conversione completata. Il file HTML è stato salvato come '{output_file}'.")
        else:
            print(f"Errore: Il formato di conversione '{args.formato_di_conversione}' non è supportato per i file PDF. Scegli 'txt' per estrarre il testo o altri formati per la conversione da testo a quel formato.")
            sys.exit(1)
    else:
        print(f"Errore: Il formato di conversione")
        sys.exit(1)

if __name__ == "__main__":
    main()

    
#####################################################################################################################

def main():
    parser = argparse.ArgumentParser(description='Converti un file in vari formati.')
    parser.add_argument('file_da_convertire', help='Il percorso del file da convertire')
    parser.add_argument('formato_di_conversione', choices=['txt', 'pdf', 'csv', 'json', 'md', 'xlsx', 'doc', 'docx', 'rtf', 'odt', 'html'], help='Il formato di conversione desiderato')

    args = parser.parse_args()

    if not os.path.isfile(args.file_da_convertire):
        print(f"Errore: Il file '{args.file_da_convertire}' non esiste.")
        sys.exit(1)

    base_name, extension = os.path.splitext(args.file_da_convertire)
    output_file = f"{base_name}.{args.formato_di_conversione}"

    count = 1
    while os.path.exists(output_file):
        output_file = f"{base_name}({count}).{args.formato_di_conversione}"
        count += 1

    text = read_text_from_file(args.file_da_convertire)

    if args.formato_di_conversione == 'pdf':
        convert_text_to_pdf(text, output_file)
        print(f"Conversione completata. Il file PDF è stato salvato come '{output_file}'.")
    elif args.formato_di_conversione == 'csv':
        convert_text_to_csv(text, output_file)
        print(f"Conversione completata. Il file CSV è stato salvato come '{output_file}'.")
    elif args.formato_di_conversione == 'json':
        convert_text_to_json(text, output_file)
        print(f"Conversione completata. Il file JSON è stato salvato come '{output_file}'.")
    elif args.formato_di_conversione == 'md':
        convert_text_to_md(text, output_file)
        print(f"Conversione completata. Il file MD è stato salvato come '{output_file}'.")
    # ADD OTHERS
    else:
        print(f"Errore: Il formato di conversione '{args.formato_di_conversione}' non è supportato. Scegli un formato tra 'pdf', 'csv', 'json', 'md', 'xlsx', 'doc', 'docx', 'rtf', 'odt', 'html'.")
        sys.exit(1)
